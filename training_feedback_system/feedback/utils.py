import openai
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
import base64
from collections import Counter
import json
from django.core.mail import EmailMessage
from django.template.loader import render_to_string
import os
import pandas as pd
import numpy as np
from collections import defaultdict
from typing import Dict, List, Optional
import warnings
warnings.filterwarnings('ignore')

# Configure OpenAI
openai.api_key = settings.OPENAI_API_KEY

def analyze_feedback_with_openai(feedback_responses):
    """
    Analyze feedback responses using OpenAI API
    """
    if not openai.api_key:
        return "OpenAI API key not configured"

    # Prepare feedback text for analysis
    feedback_texts = []
    for response in feedback_responses:
        feedback_text = f"""
        Key Learnings: {response.key_learnings}
        Missing Elements: {response.missing_elements or 'None mentioned'}
        Average Rating: {response.get_average_rating()}/5.0
        """
        feedback_texts.append(feedback_text)

    combined_feedback = "\n---\n".join(feedback_texts)

    prompt = f"""You are a qualitative feedback analyst for a training organization. You are given feedback entries from learners about a training session. Your task is to:

Analyze the feedback and provide a comprehensive overall summary that includes:
1. Overall sentiment of the feedback (positive, negative, or mixed)
2. Key strengths of the session
3. Main areas for improvement
4. Actionable recommendations for the trainer
5. A concise overall assessment in 3-5 sentences

Output ONLY a JSON object with a single key "overall_summary" containing the complete analysis as a string.

Example output format:
{{"overall_summary": "The session received mixed feedback with [analysis here]. Key strengths include [strengths]. Areas for improvement include [areas]. Recommendations: [recommendations]. Overall: [summary]."}}

Input feedback:
---
{combined_feedback}
---

Output only the JSON object, no additional text:"""

    try:
        client = openai.OpenAI(api_key=getattr(settings, 'OPENAI_API_KEY', None))
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error analyzing feedback: {str(e)}"

def create_rating_chart(feedback_responses, statement_text, statement_number):
    """
    Create a professional bar chart for rating distribution and return as base64 image string.
    """
    ratings = []
    for response in feedback_responses:
        rating_value = getattr(response, f'rating_{statement_number}')
        ratings.append(rating_value)
    
    rating_counts = Counter(ratings)
    
    # Create figure with larger size for better readability
    fig, ax = plt.subplots(figsize=(12, 6))
    
    categories = ['Strongly\nDisagree', 'Disagree', 'Neutral', 'Agree', 'Strongly\nAgree']
    values = [rating_counts.get(i, 0) for i in range(1, 6)]
    
    bars = ax.bar(categories, values, color='#4472C4', width=0.6)
    
    # Title with professional font
    ax.set_title(statement_text, fontsize=14, fontweight='bold', pad=20)
    
    # Y-axis label
    ax.set_ylabel('Number of Responses', fontsize=11, fontweight='bold')
    
    # Improve tick labels
    ax.tick_params(axis='x', labelsize=10)
    ax.tick_params(axis='y', labelsize=10)
    
    # Set y-axis to start from 0 with proper limit
    max_value = max(values) if values else 1
    ax.set_ylim(0, max_value + 1)
    
    # Add value labels on top of bars with better positioning
    for bar, value in zip(bars, values):
        if value > 0:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                    f'{int(value)}',
                    ha='center', va='bottom', fontsize=11, fontweight='bold')
    
    # Remove top and right spines for cleaner look
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Improve layout spacing
    plt.tight_layout(pad=1.5)
    
    # Save to BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    
    buf.seek(0)
    img_base64 = base64.b64encode(buf.read()).decode('utf-8')
    return f'data:image/png;base64,{img_base64}'

def generate_word_report(session, feedback_responses, ai_analysis):
    """
    Generate a professional Word report for a session, including feedback summaries and charts.
    Returns a tuple (report_path, report_filename).
    """
    report_filename = f"feedback_report_{session.id}.docx"
    report_path = os.path.join('media', 'reports', report_filename)
    os.makedirs(os.path.dirname(report_path), exist_ok=True)

    doc = Document()
    
    # Title
    doc.add_heading(f"Feedback Report: {session.session_title}", 0)
    doc.add_paragraph(f"Trainer: {session.trainer.name}")
    doc.add_paragraph(f"Date: {session.date.strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Location: {session.location}")
    doc.add_paragraph(f"Duration: {session.duration_hours} hrs")
    doc.add_paragraph("")

    # Add Professional Summary Table (Key Learnings & Missing Elements)
    doc.add_heading("Feedback Summary (Key Learnings & Missing Elements)", level=1)
    
    summary_data = None
    if ai_analysis:
        try:
            summary_data = json.loads(ai_analysis)
        except Exception:
            summary_data = None
    
    def add_bullet_section(heading, value):
        doc.add_heading(heading, level=2)
        if isinstance(value, list):
            for item in value:
                if item and str(item).strip():
                    doc.add_paragraph(str(item).lstrip('-• '), style='List Bullet')
        elif isinstance(value, str):
            lines = [line.strip('-• \n') for line in value.split('\n') if line.strip('-• \n')]
            for item in lines:
                if item:
                    doc.add_paragraph(item, style='List Bullet')
    
    if summary_data and isinstance(summary_data, dict):
        if 'overall_summary' in summary_data:
            doc.add_heading("Overall Summary", level=2)
            doc.add_paragraph(summary_data['overall_summary'])
    else:
        if ai_analysis:
            doc.add_paragraph(ai_analysis)
    
    doc.add_paragraph("")

    # Add Charts
    doc.add_heading("Feedback Charts", level=1)
    
    FEEDBACK_QUESTIONS = [
        "The training met my expectations",
        "I will be able to apply the knowledge learned",
        "The content was organized and easy to follow",
        "The trainer was knowledgeable",
        "Training was relevant to my needs",
        "Instructions were clear and understandable",
        "Length and timing of training was sufficient",
        "Overall, the session was very good"
    ]
    
    for i, question in enumerate(FEEDBACK_QUESTIONS, start=1):
        img_base64 = create_rating_chart(feedback_responses, question, i)
        if img_base64.startswith('data:image/png;base64,'):
            img_data = base64.b64decode(img_base64.split(',')[1])
            image_stream = BytesIO(img_data)
            doc.add_paragraph(question)
            doc.add_picture(image_stream, width=Inches(5.5))
            doc.add_paragraph("")

    doc.save(report_path)
    return report_path, report_filename

def send_report_email(trainer, session, report_path):
    """
    Send report email to trainer.
    """
    return True

class FeedbackAnalyzer:
    """
    A system to analyze training feedback ratings and generate analysis reports.
    """
    def __init__(self):
        self.rating_mapping = {
            1: "Strongly Disagree",
            2: "Disagree", 
            3: "Neutral",
            4: "Agree",
            5: "Strongly Agree"
        }
        self.question_mapping = {
            "Meet my expectation": "The training met my expectations",
            "Knowledge Learned": "I will be able to apply the knowledge learned",
            "Content": "The content was organized and easy to follow",
            "Trainer Knowledgeable": "The trainer was knowledgeable",
            "Training Relevancy": "Training was relevant to my needs",
            "Clear and understandable": "Instructions were clear and understandable",
            "Length Timing": "Length and timing of training was sufficient",
            "Overall": "Overall, the session was very good"
        }
